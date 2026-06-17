import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

docx_path = r"c:\Users\marcotulio\OneDrive\MBACD\Projeto\Guia_Conexao_BI.docx"

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(tblBorders)

def add_credentials_table(doc, creds):
    headers = ["Parâmetro de Conexão", "Valor de Configuração", "Descrição"]
    table = doc.add_table(rows=0, cols=len(headers))
    set_table_borders(table)
    table.autofit = False
    
    # Cabeçalho da Tabela
    hdr_row = table.add_row()
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_background(cell, "0F766E")  # Teal Escuro
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Dados da Tabela
    for r_idx, c in enumerate(creds):
        row = table.add_row()
        row_data = [c['param'], c['val'], c['desc']]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            set_cell_margins(cell)
            
            # Zebra striping
            if r_idx % 2 == 1:
                set_cell_background(cell, "F9FAFB")
                
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 65, 85)
                
            # Destacar nome do parâmetro
            if c_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(15, 23, 42)
            # Código no valor
            elif c_idx == 1:
                p.runs[0].font.name = 'Consolas'
                p.runs[0].font.color.rgb = RGBColor(15, 118, 110)
                
    # Larguras
    col_widths = [Inches(2.0), Inches(2.2), Inches(2.3)]
    for i, col in enumerate(table.columns):
        width = col_widths[i] if i < len(col_widths) else Inches(1.0)
        for cell in col.cells:
            cell.width = width
            
    doc.add_paragraph()

def generate():
    doc = Document()
    
    # Margens A4
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Fonte Padrão
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(15, 23, 42)

    # --- TÍTULO ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("GUIA DE CONEXÃO PARA FERRAMENTAS DE BI")
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 118, 110)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Integração de Power BI, Tableau e Ferramentas SQL ao Star Schema PostgreSQL\nHealth Analytics Dashboard")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    # --- 1. CREDENCIAIS ---
    p_h1 = doc.add_heading(level=1)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)
    run = p_h1.add_run("1. Parâmetros de Acesso ao Banco de Dados")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110)
    
    doc.add_paragraph(
        "Para extrair dados analíticos e modelar relatórios em ferramentas externas de BI, "
        "é necessário conectar-se diretamente ao banco de dados relacional PostgreSQL corporativo. "
        "Utilize os seguintes parâmetros de conexão:"
    )

    creds = [
        {"param": "SGDB / Driver", "val": "PostgreSQL", "desc": "Banco de dados relacional de produção."},
        {"param": "Servidor (Host)", "val": "localhost", "desc": "IP do servidor local ou de produção."},
        {"param": "Porta", "val": "5432", "desc": "Porta padrão de conexão do PostgreSQL."},
        {"param": "Banco de Dados (DB)", "val": "indicadores_db", "desc": "Nome do banco de dados do projeto."},
        {"param": "Schema Alvo", "val": "ind", "desc": "Schema onde residem as dimensões e fatos analíticos."},
        {"param": "Usuário (User)", "val": "postgres", "desc": "Usuário administrador de leitura."},
        {"param": "Senha (Password)", "val": "postgres", "desc": "Senha de acesso correspondente."}
    ]
    add_credentials_table(doc, creds)

    # --- 2. POWER BI ---
    p_h1 = doc.add_heading(level=1)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)
    run = p_h1.add_run("2. Configuração de Conexão no Microsoft Power BI")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110)

    doc.add_paragraph(
        "O Microsoft Power BI possui um conector nativo otimizado para o PostgreSQL. Siga os passos abaixo "
        "para realizar a importação de dados:"
    )

    steps_pbi = [
        "Abra o **Power BI Desktop**.",
        "Na barra de ferramentas inicial, clique em **Obter Dados (Get Data)** e selecione **Mais...**.",
        "Selecione a categoria **Banco de dados** e escolha **Banco de dados PostgreSQL**, depois clique em **Conectar**.",
        "No campo **Servidor**, digite `localhost`. No campo **Banco de dados**, digite `indicadores_db`.",
        "Na aba de credenciais de acesso, clique na aba **Banco de dados** (à esquerda), insira o usuário `postgres` e a senha `postgres`, e clique em **Conectar**.",
        "No painel de Navegação, expanda a árvore e acesse o schema **`ind`**.",
        "Marque as caixas de todas as dimensões (`dim_...`) e fatos (`fato_...`) que deseja analisar, e clique em **Transformar Dados** (altamente recomendado para configurar tipos de dados) ou **Carregar**.",
        "Na visualização de **Modelagem (Aba Relacionamentos)**, valide as chaves de ligação de 1 para muitos (1:N) de cada Dimensão para sua respectiva Fato por meio do ID (ex: `DimTempo.id_tempo` para `FatoAtendimentos.id_tempo`)."
    ]
    for step in steps_pbi:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(step)
        p.paragraph_format.space_after = Pt(3)

    # --- 3. TABLEAU ---
    p_h1 = doc.add_heading(level=1)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)
    run = p_h1.add_run("3. Configuração de Conexão no Tableau Desktop")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110)

    doc.add_paragraph(
        "Para usuários do Tableau Desktop, a conexão com o PostgreSQL é direta por meio de driver JDBC/ODBC nativo:"
    )

    steps_tab = [
        "Abra o **Tableau Desktop**.",
        "No painel esquerdo sob a seção **Conectar**, clique em **PostgreSQL** (se não estiver listado, vá em **Mais...**).",
        "Insira os dados: Servidor = `localhost`, Porta = `5432`, Banco de dados = `indicadores_db`, Nome de usuário = `postgres`, Senha = `postgres`.",
        "Clique em **Entrar**.",
        "Na aba de Fonte de Dados, clique no menu suspenso de **Esquema** e escolha o schema **`ind`**.",
        "Arraste a tabela de fatos principal que deseja analisar (ex: `FatoFinanceiro`) para a área de modelagem.",
        "Arraste as tabelas de dimensão associadas (ex: `DimConvenio`, `DimTempo`) e configure a relação (Relationship) ou junção (Join) utilizando as chaves de ID correspondentes."
    ]
    for step in steps_tab:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(step)
        p.paragraph_format.space_after = Pt(3)

    # --- 4. BOAS PRÁTICAS ---
    p_h1 = doc.add_heading(level=1)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)
    run = p_h1.add_run("4. Melhores Práticas de Performance e Modelagem")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110)

    bp = [
        "**Modo Import vs DirectQuery:** Para a maioria dos dashboards corporativos hospitalares, prefira o modo de **Importação (Import)** com agendamento de atualização (ex: 2 ou 4 vezes ao dia). O modo DirectQuery gera consultas SQL diretas no banco relacional a cada clique do usuário, o que pode impactar a performance do banco transacional sob muita carga.",
        "**Redução de Colunas:** Importe apenas os campos estritamente necessários para a análise. Campos textuais muito longos ou IDs de auditoria que não são chaves de relacionamento aumentam desnecessariamente o tamanho do arquivo local (`.pbix`).",
        "**Formatação de Datas:** Utilize o campo `id_tempo` (ex: `20250701`) como chave primária de ligação física, mas configure a exibição de eixos cronológicos usando o campo `data_registro` formatado ou crie uma tabela de calendário dTempo interna no Power BI para suporte a inteligência de tempo (Time Intelligence).",
        "**Esquema Estrela Puro:** Evite mesclar tabelas Fato com tabelas Dimensão no Power Query (denormalização excessiva). Manter o Star Schema preservado melhora consideravelmente o desempenho dos cálculos DAX e a legibilidade do modelo de dados."
    ]
    for b in bp:
        p = doc.add_paragraph(style='List Bullet')
        # Formatar negrito
        parts = b.split("**")
        if len(parts) >= 3:
            p.add_run(parts[1]).font.bold = True
            p.add_run(parts[2])
        else:
            p.add_run(b)
        p.paragraph_format.space_after = Pt(4)

    # Salvar
    doc.save(docx_path)
    print(f"Sucesso! Guia de Conexão BI gerado em: {docx_path}")

if __name__ == '__main__':
    generate()
