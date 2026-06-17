import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

docx_path = r"c:\Users\marcotulio\OneDrive\MBACD\Projeto\Matriz_Metricas_KPIs_Hospitalar.docx"

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(tblBorders)

def add_kpi_table(doc, kpis):
    headers = ["Indicador (KPI)", "Fórmula de Cálculo", "Origem dos Dados", "Objetivo Estratégico"]
    table = doc.add_table(rows=0, cols=len(headers))
    set_table_borders(table)
    table.autofit = False
    
    # Cabeçalho da Tabela
    hdr_row = table.add_row()
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_background(cell, "0F766E")  # Teal Escuro Corporativo (#0f766e)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Dados da Tabela
    for r_idx, kpi in enumerate(kpis):
        row = table.add_row()
        row_data = [kpi['nome'], kpi['formula'], kpi['origem'], kpi['objetivo']]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            set_cell_margins(cell)
            
            # Zebra striping
            if r_idx % 2 == 1:
                set_cell_background(cell, "F0FDFA") # Light Teal
                
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
                
            # Destacar nome do indicador em negrito
            if c_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(15, 23, 42)
                
    # Larguras das Colunas
    col_widths = [Inches(1.8), Inches(2.2), Inches(1.3), Inches(2.2)]
    for i, col in enumerate(table.columns):
        width = col_widths[i] if i < len(col_widths) else Inches(1.0)
        for cell in col.cells:
            cell.width = width
            
    doc.add_paragraph() # Espaço

def generate():
    doc = Document()
    
    # Margens A4 padrão
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configuração de Fonte Padrão
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(15, 23, 42)

    # --- CAPA OU TÍTULO PRINCIPAL ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("MATRIZ DE DIRETRIZES, MÉTRICAS E KPIs")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 118, 110) # Teal

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Manual de Fórmulas e Regras de Negócio para Apresentação à Diretoria Executiva\nHealth Analytics Dashboard - BI Hospitalar")
    run_sub.font.size = Pt(11.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    # --- INTRODUÇÃO EXECUTIVA ---
    p_h1 = doc.add_heading(level=1)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)
    run = p_h1.add_run("1. Visão Geral Estratégica")
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110)
    
    doc.add_paragraph(
        "A gestão baseada em dados é o pilar fundamental para garantir a eficiência operacional e a "
        "sustentabilidade financeira das instituições de saúde modernas. Este manual consolida a Matriz de KPIs "
        "e Fórmulas de Negócio utilizadas no **Health Analytics Dashboard**."
    )
    doc.add_paragraph(
        "O objetivo desta matriz é fornecer à Diretoria uma visão padronizada de como os indicadores de performance "
        "são calculados, as fontes originais dos dados (modelagem Star Schema) e a justificativa estratégica de "
        "cada métrica nas tomadas de decisão de alto nível."
    )

    # --- DETALHAMENTO DOS MÓDULOS ---
    modulos = [
        {
            "titulo": "2. Módulo de Pacientes (Experiência e Fluxo)",
            "desc": "Focado no monitoramento da capacidade assistencial, satisfação e qualidade clínica da jornada do paciente no hospital.",
            "kpis": [
                {
                    "nome": "Taxa de Ocupação",
                    "formula": "(Leitos Ocupados / Capacidade Máxima de Leitos) * 100",
                    "origem": "DimUnidade, FatoAtendimentos",
                    "objetivo": "Gestão de capacidade física e leitos disponíveis, evitando superlotação ou ociosidade."
                },
                {
                    "nome": "Tempo Médio de Permanência (TMP)",
                    "formula": "Soma(Tempo de Permanência em Dias) / Total de Pacientes com Alta",
                    "origem": "FatoAtendimentos",
                    "objetivo": "Identificar gargalos nos fluxos de alta e aumentar o giro de leitos sem prejudicar o desfecho."
                },
                {
                    "nome": "Net Promoter Score (NPS)",
                    "formula": "% Pacientes Promotores (NPS 9-10) - % Pacientes Detratores (NPS 0-6)",
                    "origem": "DimPaciente",
                    "objetivo": "Avaliar qualitativamente a experiência global do paciente e a qualidade do atendimento percebido."
                },
                {
                    "nome": "Taxa de Reinternação (30 dias)",
                    "formula": "(Pacientes Reinternados em até 30 dias / Total de Altas) * 100",
                    "origem": "DimPaciente, FatoAtendimentos",
                    "objetivo": "Principal indicador de qualidade clínica e resolutividade. Evita altas precoces que geram retorno."
                }
            ]
        },
        {
            "titulo": "3. Módulo de Medicamentos (Suprimentos e Segurança)",
            "desc": "Monitora os custos com farmácia, a rotatividade de insumos críticos e a mitigação de incidentes assistenciais.",
            "kpis": [
                {
                    "nome": "Giro de Estoque",
                    "formula": "Total de Medicamentos Dispensados no Mês / Estoque Médio Mensal",
                    "origem": "DimMedicamento, FatoEstoque",
                    "objetivo": "Otimizar capital de giro. Evita o vencimento de fármacos caros e minimiza custos de armazenamento."
                },
                {
                    "nome": "Custo de Farmácia por Paciente/Dia",
                    "formula": "Custo Total de Medicamentos Dispensados / Soma(Dias de Internação)",
                    "origem": "DimMedicamento, FatoEstoque, FatoAtendimentos",
                    "objetivo": "Controlar a sinistralidade e desvios de consumo terapêutico por paciente internado."
                },
                {
                    "nome": "Índice de Erros de Medicação",
                    "formula": "(Ocorrências com erro de medicação / Total de Atendimentos) * 1000",
                    "origem": "FatoErrosMedicao, FatoAtendimentos",
                    "objetivo": "Indicador vital do Núcleo de Segurança do Paciente (NSP). Visa mitigar falhas assistenciais graves."
                },
                {
                    "nome": "Taxa de Ruptura de Estoque",
                    "formula": "(Medicamentos Críticos Faltantes / Total Solicitado) * 100",
                    "origem": "DimMedicamento, FatoEstoque",
                    "objetivo": "Medir a eficiência do setor de compras. Evita a descontinuidade de tratamentos vitais."
                }
            ]
        },
        {
            "titulo": "4. Módulo de Procedimentos (Produção e Bloco Cirúrgico)",
            "desc": "Avalia o desempenho e aproveitamento dos recursos mais caros da estrutura hospitalar: salas cirúrgicas e equipamentos de imagem.",
            "kpis": [
                {
                    "nome": "Taxa de Suspensão de Cirurgias",
                    "formula": "(Procedimentos Cancelados / Total de Agendamentos Cirúrgicos) * 100",
                    "origem": "FatoProcedimentos",
                    "objetivo": "Reduzir ociosidade cirúrgica, retrabalhos administrativos e insatisfação de pacientes."
                },
                {
                    "nome": "Tempo de Giro de Sala (Turnover)",
                    "formula": "Tempo Médio (Minutos de Preparo + Minutos de Higienização)",
                    "origem": "FatoProcedimentos",
                    "objetivo": "Aumentar a produtividade do centro cirúrgico otimizando os tempos mortos entre cirurgias."
                },
                {
                    "nome": "Produtividade por Equipamento",
                    "formula": "Soma(Tempo de Execução Realizado) / Horas Totais de Capacidade Ativa",
                    "origem": "DimEquipamento, FatoProcedimentos",
                    "objetivo": "Avaliar o retorno sobre investimento (ROI) de aparelhos de alto custo (Ressonância, Tomografia)."
                },
                {
                    "nome": "Taxa de Conversão Cirúrgica",
                    "formula": "(Cirurgias Realizadas / Consultas Especializadas Realizadas) * 100",
                    "origem": "DimTipoProcedimento, FatoProcedimentos",
                    "objetivo": "Avaliar a taxa de resolutividade cirúrgica a partir dos atendimentos ambulatoriais."
                }
            ]
        },
        {
            "titulo": "5. Módulo Hospitalar (Infraestrutura e Hotelaria)",
            "desc": "Acompanha os gastos indiretos, sustentabilidade ambiental do hospital e segurança ambiental contra infecções.",
            "kpis": [
                {
                    "nome": "Taxa de Infecção Hospitalar (IRAS)",
                    "formula": "(Eventos de Infecção Ativos / Total de Pacientes-Dia no Mês) * 1000",
                    "origem": "FatoInfraestrutura, FatoAtendimentos",
                    "objetivo": "Indicador de segurança epidemiológica (CCIH). Visa zerar infecções associadas à internação."
                },
                {
                    "nome": "Intervalo de Substituição de Leito",
                    "formula": "Média(data_hora_liberacao_leito - data_hora_saida_paciente)",
                    "origem": "FatoHigienizacao",
                    "objetivo": "Medir a performance do time de hotelaria para limpeza e setup de leitos para novos pacientes."
                },
                {
                    "nome": "Consumo de Recursos por Leito/Dia",
                    "formula": "Consumo Mensal (Água ou Luz) / Soma(Diárias de Leito Ocupado)",
                    "origem": "FatoInfraestrutura, FatoAtendimentos",
                    "objetivo": "Métricas de sustentabilidade (ESG) e controle de custos fixos de instalações prediais."
                },
                {
                    "nome": "Densidade de Recursos Humanos",
                    "formula": "Total de Enfermeiros e Assistentes / Total de Leitos Ativos na Ala",
                    "origem": "FatoInfraestrutura, DimUnidade",
                    "objetivo": "Avaliar a proporção de colaboradores por paciente para garantir o padrão de segurança assistencial."
                }
            ]
        },
        {
            "titulo": "6. Módulo Corpo Clínico (Médicos)",
            "desc": "Mede o engajamento dos médicos com as boas práticas clínicas, prontidão de registros e assiduidade.",
            "kpis": [
                {
                    "nome": "Adesão a Protocolos Clínicos",
                    "formula": "(Condutas Médicas em Conformidade / Total de Casos Avaliados) * 100",
                    "origem": "DimProtocolo, FatoDesempenhoClinico",
                    "objetivo": "Garantir a Medicina Baseada em Evidências, reduzindo a variabilidade clínica e desfechos desfavoráveis."
                },
                {
                    "nome": "Tempo para Fechamento de Prontuário",
                    "formula": "Média(Tempo decorrido entre a Alta e a Assinatura do Prontuário)",
                    "origem": "FatoDesempenhoClinico",
                    "objetivo": "Garantir o faturamento ágil das contas. Prontuários abertos impedem o envio de cobranças às operadoras."
                },
                {
                    "nome": "Taxa de Absenteísmo Médico",
                    "formula": "(Plantões com Falta ou Atraso / Total de Plantões Agendados) * 100",
                    "origem": "DimMedico, FatoEscalaMedica",
                    "objetivo": "Monitorar a assiduidade e a cobertura assistencial para garantir o suporte de plantão de emergência."
                },
                {
                    "nome": "Carga de Atendimento por Especialidade",
                    "formula": "Total de Atendimentos Realizados no Mês / Quantidade de Médicos Ativos",
                    "origem": "DimMedico, FatoAtendimentos",
                    "objetivo": "Balanceamento da carga de trabalho médico por especialidade para evitar sobrecarga (burnout)."
                }
            ]
        },
        {
            "titulo": "7. Módulo Financeiro e Recursos Orçamentários",
            "desc": "Mede a rentabilidade operacional, gargalos de glosa e eficiência do ciclo de faturamento hospitalar.",
            "kpis": [
                {
                    "nome": "EBITDA Hospitalar",
                    "formula": "Receita Bruta Faturada - Custos Operacionais Totais",
                    "origem": "FatoFinanceiro",
                    "objetivo": "Principal métrica de lucro operacional antes dos efeitos fiscais e financeiros. Mede a eficiência do hospital."
                },
                {
                    "nome": "Ticket Médio por Convênio",
                    "formula": "Faturamento Bruto Total do Convênio / Total de Atendimentos do Convênio",
                    "origem": "DimConvenio, FatoFinanceiro",
                    "objetivo": "Análise de rentabilidade de cada operadora de plano de saúde parceira para negociações de tabela."
                },
                {
                    "nome": "Índice de Glosa Inicial",
                    "formula": "(Valor Total das Cobranças Recusadas / Receita Bruta Faturada) * 100",
                    "origem": "DimConvenio, FatoFinanceiro",
                    "objetivo": "Identificar falhas de preenchimento e autorização de contas médicas junto às operadoras."
                },
                {
                    "nome": "Taxa de Recuperação de Glosa",
                    "formula": "(Valor de Glosas Recuperadas após Recurso / Valor de Glosa Inicial) * 100",
                    "origem": "FatoFinanceiro",
                    "objetivo": "Avaliar o desempenho do time de faturamento e auditoria de recursos de glosa administrativa."
                },
                {
                    "nome": "Prazo Médio de Recebimento (PMR)",
                    "formula": "Média(Data de Depósito Real do Convênio - Data de Emissão da Nota)",
                    "origem": "DimConvenio, FatoFinanceiro",
                    "objetivo": "Controlar o fluxo de caixa. Prazos elevados exigem maior necessidade de capital de giro."
                }
            ]
        }
    ]

    for mod in modulos:
        p_h1 = doc.add_heading(level=1)
        p_h1.paragraph_format.space_before = Pt(18)
        p_h1.paragraph_format.space_after = Pt(4)
        run = p_h1.add_run(mod["titulo"])
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 118, 110)
        
        doc.add_paragraph(mod["desc"]).paragraph_format.space_after = Pt(8)
        add_kpi_table(doc, mod["kpis"])

    # Salvar
    doc.save(docx_path)
    print(f"Sucesso! Matriz de KPIs gerada em: {docx_path}")

if __name__ == '__main__':
    generate()
