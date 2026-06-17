import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

docx_path = r"c:\Users\marcotulio\OneDrive\MBACD\Projeto\Especificacao_API_REST_Swagger.docx"

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(tblBorders)

def add_styled_table(doc, headers, data, col_widths):
    table = doc.add_table(rows=0, cols=len(headers))
    set_table_borders(table)
    table.autofit = False
    
    # Adicionar cabeçalho
    hdr_row = table.add_row()
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_background(cell, "1E3A8A")  # Azul Escuro Corporate (#1e3a8a)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Adicionar dados
    for r_idx, row_data in enumerate(data):
        row = table.add_row()
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
            # Estilo zebra
            if r_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC") # Slate 50
                
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
                
    # Definir larguras
    for i, col in enumerate(table.columns):
        width = col_widths[i] if i < len(col_widths) else Inches(1.0)
        for cell in col.cells:
            cell.width = width
            
    doc.add_paragraph() # Espaço

def generate():
    doc = Document()
    
    # Margens A4
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Fonte Padrão
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(15, 23, 42)

    # 1. TÍTULO DO DOCUMENTO
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("ESPECIFICAÇÃO DA API REST & DOCUMENTAÇÃO SWAGGER")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 41, 59) # Slate 800

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Health Analytics Dashboard (BI Hospitalar)\nVersão da API: 1.0.0 | Padrão OpenAPI 3.0")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
    doc.add_paragraph()

    # 2. INTRODUÇÃO
    p_h1 = doc.add_heading(level=1)
    run = p_h1.add_run("1. Introdução e Visão Geral")
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "A API REST do Health Analytics Dashboard foi construída utilizando Django e Django REST Framework "
        "para atuar como o barramento analítico da instituição. Ela expõe endpoints estruturados para "
        "carga de dados em lote, consulta operacional e alimentação de BI Tools externas (como Power BI ou Tableau)."
    )

    doc.add_paragraph(
        "A documentação interativa da API está baseada no padrão de especificação técnica **OpenAPI 3.0** "
        "por meio da biblioteca `drf-spectacular`. As interfaces interativas expostas na aplicação permitem testar "
        "as chamadas diretamente no navegador."
    )

    # 3. INTERFACES INTERATIVAS (SWAGGER E REDOC)
    p_h1 = doc.add_heading(level=1)
    run = p_h1.add_run("2. Endpoints da Documentação Swagger e Redoc")
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "Quando o servidor local estiver rodando (porta 8000), os seguintes caminhos estarão disponíveis "
        "para acesso direto via navegador:"
    )

    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run("Esquema OpenAPI (JSON/YAML): ")
    run_bold.font.bold = True
    p.add_run("http://localhost:8000/api/schema/")
    
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run("Swagger UI (Visualizador Interativo): ")
    run_bold.font.bold = True
    p.add_run("http://localhost:8000/api/schema/swagger-ui/")
    
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run("ReDoc UI (Documentação Técnica Avançada): ")
    run_bold.font.bold = True
    p.add_run("http://localhost:8000/api/schema/redoc/")

    # 4. SEGURANÇA E AUTENTICAÇÃO
    p_h1 = doc.add_heading(level=1)
    run = p_h1.add_run("3. Segurança e Autenticação por Token")
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "A segurança da API é gerida pelo mecanismo de autenticação por tokens (`TokenAuthentication` do DRF). "
        "Qualquer chamada para sincronização de dados ou inserção deve conter a chave de autorização inserida no "
        "cabeçalho HTTP da requisição."
    )

    doc.add_paragraph(
        "Formato do cabeçalho de autorização:"
    )
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.5)
    run_code = p_code.add_run("Authorization: Token 618a1ecc975393339a8f8c1a35c779cc142e7654")
    run_code.font.name = 'Consolas'
    run_code.font.size = Pt(9.5)
    run_code.font.color.rgb = RGBColor(15, 118, 110) # Teal

    # 5. TABELA DE ENDPOINTS DE INTEGRAÇÃO
    p_h1 = doc.add_heading(level=1)
    run = p_h1.add_run("4. Endpoints Globais de Ingestão e Status")
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)

    headers_int = ["Método", "Endpoint", "Content-Type", "Descrição"]
    data_int = [
        ["POST", "/api/integracoes/sync/<table_id>/", "application/json ou text/csv", "Sincroniza registros em lote (Upsert) de qualquer tabela do Star Schema. Requer delimitador ';' para arquivos CSV."],
        ["GET", "/api/integracoes/status/", "application/json", "Retorna a listagem histórica e o timestamp das últimas cargas de cada tabela fato ou dimensão."]
    ]
    add_styled_table(doc, headers_int, data_int, [Inches(1.0), Inches(2.2), Inches(1.5), Inches(2.8)])

    # 6. TABELA DE ENDPOINTS DO STAR SCHEMA
    p_h1 = doc.add_heading(level=1)
    run = p_h1.add_run("5. Endpoints de Entidades do Star Schema")
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "A API disponibiliza endpoints CRUD completos (Create, Read, Update, Delete) para todas as 18 tabelas "
        "do Star Schema. A tabela abaixo resume os caminhos base para chamadas operacionais:"
    )

    headers_schema = ["Recurso Analítico", "Caminho Base (GET, POST)", "Parâmetro ID (GET, PUT, DELETE)"]
    data_schema = [
        ["Dimissão: Tempo", "/api/tempo/", "/api/tempo/<id_tempo>/"],
        ["Dimissão: Unidades", "/api/unidades/", "/api/unidades/<id_unidade>/"],
        ["Dimissão: Pacientes", "/api/pacientes/", "/api/pacientes/<id_paciente>/"],
        ["Dimissão: Medicamentos", "/api/medicamentos/", "/api/medicamentos/<id_medicamento>/"],
        ["Dimissão: Tipos Procedimento", "/api/tipos-procedimento/", "/api/tipos-procedimento/<id_tipo_procedimento>/"],
        ["Dimissão: Equipamentos", "/api/equipamentos/", "/api/equipamentos/<id_equipamento>/"],
        ["Dimissão: Protocolos", "/api/protocolos/", "/api/protocolos/<id_protocolo>/"],
        ["Dimissão: Médicos", "/api/medicos/", "/api/medicos/<id_medico>/"],
        ["Dimissão: Convênios", "/api/convenios/", "/api/convenios/<id_convenio>/"],
        ["Fato: Atendimentos", "/api/atendimentos/", "/api/atendimentos/<id_atendimento>/"],
        ["Fato: Estoque", "/api/estoque/", "/api/estoque/<id_movimentacao>/"],
        ["Fato: Erros Medicação", "/api/erros-medicao/", "/api/erros-medicao/<id_evento>/"],
        ["Fato: Procedimentos", "/api/procedimentos/", "/api/procedimentos/<id_procedimento_instancia>/"],
        ["Fato: Infraestrutura", "/api/infraestrutura/", "/api/infraestrutura/<id_registro>/"],
        ["Fato: Higienização", "/api/higienizacao/", "/api/higienizacao/<id_higienizacao>/"],
        ["Fato: Desempenho Clínico", "/api/desempenho-clinico/", "/api/desempenho-clinico/<id_registro>/"],
        ["Fato: Escala Médica", "/api/escala-medica/", "/api/escala-medica/<id_escala>/"],
        ["Fato: Financeiro", "/api/financeiro/", "/api/financeiro/<id_transacao>/"]
    ]
    add_styled_table(doc, headers_schema, data_schema, [Inches(2.3), Inches(2.2), Inches(2.5)])

    # 7. CÓDIGOS DE RESPOSTAS E EXEMPLOS
    p_h1 = doc.add_heading(level=1)
    run = p_h1.add_run("6. Respostas HTTP e Formatos de Erro")
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)
    p_h1.paragraph_format.space_before = Pt(18)
    p_h1.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "A API responde com códigos HTTP padrão de acordo com o resultado da transação:"
    )

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("200 OK: ").font.bold = True
    p.add_run("Requisição concluída com sucesso (retorno de dados ou confirmação de sincronismo).")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("201 Created: ").font.bold = True
    p.add_run("Novo registro criado com sucesso por meio de inserção via POST.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("400 Bad Request: ").font.bold = True
    p.add_run("Corpo de requisição inválido, falha de validação lógica (ex: horas_atraso preenchidas para médico presente) ou erro de parsing de arquivo CSV.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("401 Unauthorized: ").font.bold = True
    p.add_run("Cabeçalho de autorização ausente ou token inválido.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("404 Not Found: ").font.bold = True
    p.add_run("Recurso ou tabela de destino não localizada.")

    doc.add_paragraph()
    doc.add_paragraph("Exemplo de resposta para erro de validação (400 Bad Request):")
    p_err = doc.add_paragraph()
    p_err.paragraph_format.left_indent = Inches(0.5)
    p_err.add_run(
        "{\n"
        '  "erro": "Falha ao processar CSV: campo id_tempo obrigatório ausente"\n'
        "}"
    ).font.name = 'Consolas'

    # Salvar
    doc.save(docx_path)
    print(f"Sucesso! Documento DOCX da API gerado em: {docx_path}")

if __name__ == '__main__':
    generate()
